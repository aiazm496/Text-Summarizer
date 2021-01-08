#!/usr/bin/env python
# coding: utf-8

# In[13]:


import nltk
from nltk.corpus import stopwords
#from nltk.cluster.util import cosine_distance  #for similarity b/w sentences
#import numpy as np
#import networkx as nx  #for similarity graphs


# In[14]:


import os


# In[42]:


def read_feedback(file):
    if os.path.exists(file):
        file_name = os.path.abspath(file)
        with open(file_name,'r',encoding='utf-8') as fob:
            return fob.read().replace('\n','')
    else:
        print("Please check file path!")


# In[16]:


from gensim.summarization.summarizer import summarize


# In[17]:


from gensim.summarization import keywords


# In[18]:


from textblob import TextBlob


# In[19]:


file_path = input("Please enter file path: ")
file_reader = read_feedback(file_path)


# In[20]:


file_reader


# In[21]:


file_blob = TextBlob(file_reader)


# In[22]:


def translator(blob):
    if blob.detect_language()!='en':
        return blob.translate(to='en')
    else:
        return blob
            


# In[23]:


en_blob = translator(file_blob)


# In[24]:


en_blob.sentiment.polarity


# In[25]:


en_blob.sentiment.subjectivity


# In[26]:


en_file = en_blob.string


# In[27]:


file_summary = summarize(en_file,ratio = 0.3)   #summary is 30% of the original text.
#print(editorial_summary)


# In[28]:


#print(keywords(file_summary,lemmatize=True))


# In[29]:


import docx


# In[30]:


from docx import Document


# In[31]:


response = Document()


# In[32]:


response.add_heading('Summary', 0)


# In[33]:


response.add_paragraph(file_summary)


# In[34]:


#%pwd


# In[35]:


path_to_save = input("Please enter location where you want to save the file: ")


# In[36]:


os.chdir(path_to_save)  


# In[37]:


name_of_summary_file = input("Please enter summary file name")


# In[38]:


response.save('{}.docx'.format(name_of_summary_file))


# In[ ]:




